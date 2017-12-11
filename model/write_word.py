from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


class WriteWord():
    def __init__(self):
        self.document = Document()

    def add_head(self, head_name, head_level):
        # 增加不同级别的标题
        self.document.add_heading(head_name, head_level)

    def add_text(self, text, font_size=24, font='Consolas'):
        # 添加文本
        paragraph = self.document.add_paragraph(text)
        # 设置字号
        run_set_font_size = paragraph.add_run(text)
        run_set_font_size.font.size = Pt(font_size)
        # 设置字体
        run = paragraph.add_run('Set Font,')
        run.font.name = font
        # 设置中文字体
        run = paragraph.add_run(u'设置中文字体，')
        run.font.name = u'宋体'
        r = run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
